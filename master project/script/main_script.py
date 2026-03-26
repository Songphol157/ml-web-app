# Master Project - Main Script
# Machine Learning for Survival Analysis in AML
# Python version 3.12
# 03/23/2026
# by Steve Tungjitviboonkun

# Required packages are listed in requirements.txt
# Install them using: pip install -r requirements.txt

#Master Project/
#├── data/
#│   └── data_clinical.csv
#│   └── data_clinical_template.csv
#├── script/
#│   └── main_script.py
#│   └── ext_val_script.py
#├── output/
#│   ├── models/
#│   └── figures/

# Pipeline:
# 1. Load and clean data
# 2. Survival modeling with Cox, RSF, GBM (cross-validation)
# 3. Binary modeling for 12-month mortality
# 4. Bootstrap robustness analysis
# 5. Export results to DOCX

import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import joblib
try:
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
except NameError:
    BASE_DIR = os.getcwd()

# --- Paths ---
DATA_PATH = os.path.join(BASE_DIR, "data", "data_clinical.csv")

OUTPUT_DIR = os.path.join(BASE_DIR, "output")
MODELS_DIR = os.path.join(OUTPUT_DIR, "models")
FIGURES_DIR = os.path.join(OUTPUT_DIR, "figures")

# --- Create directories ---
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(MODELS_DIR, exist_ok=True)
os.makedirs(FIGURES_DIR, exist_ok=True)

# --- Validate data path ---
if not os.path.exists(DATA_PATH):
    raise FileNotFoundError(f"Data file not found at: {DATA_PATH}")

df = pd.read_csv(DATA_PATH)
print("Data loaded. First 5 rows:")
print(df.head())

# Summary Statistics
n_total = len(df)
print('Total patients in external dataset:', n_total)
n_missing_death_info = df['death'].isna().sum() # missing death information (cannot be included in survival analysis)
print('Patients with missing death information:', n_missing_death_info)
df = df[df['death'].notna()]
# drop if os_months < 0 since this is likely an error and we can't use them for survival analysis
n_negative_os = (df['os_months'] < 0).sum()
print('Patients with negative OS months (invalid):', n_negative_os)
df = df[df['os_months'] >= 0]
print('Patients included in survival analysis:', len(df))
# Exclude from binary model if censored <12 months (outcome unknown)
uncertain_mask = ((df['death'] == 'no') | (df['death'] == 'No') | (df['death'] == 'NO')) & (df['os_months'] < 12)
print('Patients excluded from binary model due to uncertain 12-month outcome:', uncertain_mask.sum())

# state the predictors, outcomes we will use
predictors = [
    'age_at_diagnosis', 'sex', 'eln2017mode', 'denovo_cat', 'chromosome_cat',
    'flt3_itd_cat', 'npm1_cat', 'runx1_cat', 'asxl1_cat', 'tp53_cat'
    ]
outcomes = ['os_months', 'death']

# Keep only predictors or outcomes columns
df = df.loc[:, predictors + outcomes]
print(df.head())

# Step 2: Data Cleaning - Lower-case text columns
for col in [c for c in predictors + ['death'] if c != 'age_at_diagnosis']:
    if col in df.columns:
        # Convert to string, then lowercase, then strip whitespace
        df[col] = df[col].astype(str).str.lower().str.strip()
        df = df.replace('nan', np.nan)

# Data Cleaning - Fix 'eln2017mode' column: change variable to Null if not really known
df.loc[df['eln2017mode'] == 'favorableorintermediate', 'eln2017mode'] = 'intermediate'
df.loc[df['eln2017mode'] == 'intermediateoradverse', 'eln2017mode'] = 'adverse'
df.loc[df['eln2017mode'] == 'missingkaryo', 'eln2017mode'] = np.nan
df.loc[df['eln2017mode'] == 'missingmutations', 'eln2017mode'] = np.nan
df.loc[df['eln2017mode'] == 'noninitial', 'eln2017mode'] = np.nan

# Fix 0.00-month survival times so the models don't crash
df['os_months'] = df['os_months'].replace(0, 0.01)

# Step 3: Summary table for age_at_diagnosis and os_months from full data (df table)
summary_vars = ['age_at_diagnosis', 'os_months']
summary_table = []
for var in summary_vars:
    col = df[var]
    mean = round(col.mean(), 2)
    median = round(col.median(), 2)
    min_ = round(col.min(), 2)
    max_ = round(col.max(), 2)
    n_missing = col.isna().sum()
    pct_missing = round(100 * n_missing / len(col), 2)
    summary_table.append({
        'Variable': var,
        'Mean': mean,
        'Median': median,
        'Min': min_,
        'Max': max_,
        'Missing (n)': n_missing,
        'Missing (%)': pct_missing
    })
summary_df = pd.DataFrame(summary_table)
print("\nSummary table for age_at_diagnosis and os_months:")
print(summary_df)

# Step 4: Table for gene mutation columns (positive/negative/missing)
gene_cols = ['flt3_itd_cat', 'npm1_cat', 'runx1_cat', 'asxl1_cat', 'tp53_cat']
gene_summary = []
for col in gene_cols:
    value_counts = df[col].value_counts(dropna=False)
    pos = value_counts.get('positive', 0)
    neg = value_counts.get('negative', 0)
    missing = value_counts.get(np.nan, 0) if np.nan in value_counts else df[col].isna().sum()
    total = len(df)
    gene_summary.append({
        'Gene': col,
        'Positive (n)': pos,
        'Positive (%)': round(100 * pos / total, 2),
        'Negative (n)': neg,
        'Negative (%)': round(100 * neg / total, 2),
        'Missing (n)': missing,
        'Missing (%)': round(100 * missing / total, 2)
    })
gene_summary_df = pd.DataFrame(gene_summary)
print("\nGene mutation summary table:")
print(gene_summary_df)

# Step 5: Table for sex (male/female/missing)
sex_counts = df['sex'].value_counts(dropna=False)
total = len(df)
male = sex_counts.get('male', 0)
female = sex_counts.get('female', 0)
missing = df['sex'].isna().sum()
total = len(df)
sex_summary = pd.DataFrame({
    'Sex': ['male', 'female', 'missing', 'total'],
    'Count': [male, female, missing, total],
    'Percent': [round(100 * male / total, 2), round(100 * female / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\nSex summary table:")
print(sex_summary)

# Step 6: Table for denovo_cat (True/False/Missing)
denovo_counts = df['denovo_cat'].value_counts(dropna=False)
total = len(df)
true_ = denovo_counts.get('true', 0)
false_ = denovo_counts.get('false', 0)
missing = df['denovo_cat'].isna().sum()
denovo_summary = pd.DataFrame({
    'denovo_cat': ['true', 'false', 'missing', 'total'],
    'Count': [true_, false_, missing, total],
    'Percent': [round(100 * true_ / total, 2), round(100 * false_ / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\ndenovo_cat summary table:")
print(denovo_summary)

# Step 7: Table for eln2017mode (favorable/intermediate/adverse/missing)
eln_counts = df['eln2017mode'].value_counts(dropna=False)
total = len(df)
fav = eln_counts.get('favorable', 0)
intm = eln_counts.get('intermediate', 0)
advs = eln_counts.get('adverse', 0)
missing = df['eln2017mode'].isna().sum()
eln_summary = pd.DataFrame({
    'ELN class': ['favorable', 'intermediate', 'adverse', 'missing', 'total'],
    'Count': [fav, intm, advs, missing, total],
    'Percent': [round(100 * fav / total, 2), round(100 * intm / total, 2), round(100 * advs / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\nELN 2017 classification summary table:")
print(eln_summary)

# Step 8: Table for chromosome_cat (normal/complex/missing)
chrom_counts = df['chromosome_cat'].value_counts(dropna=False)
total = len(df)
normal = chrom_counts.get('normal', 0)
complex_ = chrom_counts.get('complex', 0)
missing = df['chromosome_cat'].isna().sum()
chrom_summary = pd.DataFrame({
    'chromosome cat': ['normal', 'complex', 'missing', 'total'],
    'Count': [normal, complex_, missing, total],
    'Percent': [round(100 * normal / total, 2), round(100 * complex_ / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\nchromosome category summary table:")
print(chrom_summary)

# Use pd.get_dummies for one-hot encoding of categorical variables, with drop_first=True to avoid multicollinearity
cat_cols = [
    'sex', 'chromosome_cat', 'denovo_cat',
    'flt3_itd_cat', 'npm1_cat', 'runx1_cat',
    'asxl1_cat', 'tp53_cat'
]

# Map eln2017mode to numeric for modeling (favorable=0, intermediate=1, adverse=2)
df['eln2017mode'] = df['eln2017mode'].map({'favorable': 0, 'intermediate': 1, 'adverse': 2})
df['death'] = df['death'].map({'no': 0, 'yes': 1})
print(df.dtypes)
print(df.groupby('eln2017mode')['os_months'].mean()) # check mean os_months by eln2017mode

# Step 9: Exclude censoring within 12 months from this binary model since we don't know if they died within 12 months or not)
# Create a 'df_binary' dataframe that EXCLUDES them (using the ~ symbol)
df_binary = df[~uncertain_mask].copy() # create a separate dataframe for the binary model that excludes uncertain cases
print('Patients included in binary analysis:', len(df_binary))
df_binary['12_mo_death']=(df_binary['os_months'] < 12).astype(int)
df_binary = df_binary[predictors + ['12_mo_death']] # keep only predictors and binary outcome for the binary model
print(df_binary.head())
print(df_binary.dtypes)

# Step 10: Use Repeated Stratified K-Fold (5 folds x 10 repeats) using the FULL dataset (df data)
from sklearn.model_selection import RepeatedStratifiedKFold
from lifelines import CoxPHFitter
from lifelines.utils import concordance_index
from sksurv.ensemble import RandomSurvivalForest, GradientBoostingSurvivalAnalysis
from sksurv.util import Surv
from sklearn.metrics import roc_curve, auc, brier_score_loss
from sksurv.metrics import cumulative_dynamic_auc
from sklearn.impute import SimpleImputer

mode_imputer = SimpleImputer(strategy='most_frequent')

# Use only predictor columns for X and os_months and death for y
X = df[predictors] # predictors
y_time = df['os_months'] # time to event
y_event = df['death'] # event indicator (1 if death, 0 if censored)
# Delete the rows with missing values in any of the outcomes
mask = y_time.notna() & y_event.notna()
X = X[mask]
y_time = y_time[mask]
y_event = y_event[mask]

# Repeated stratified k-fold on the full dataset
rskf = RepeatedStratifiedKFold(n_splits=5, n_repeats=10, random_state=10)
print("\nUsing Repeated Stratified K-Fold (5 folds x 10 repeats) on the FULL dataset.")

# Lists to collect Cross-validation metrics in Cox, RSF, GBM survival models=======================================
cindex_eln_cv, cindex_cox_cv, cindex_rsf_cv, cindex_gbm_cv = [], [], [], []
auroc_eln_cv, auroc_cox_cv, auroc_rsf_cv, auroc_gbm_cv = [], [], [], []

times = np.arange(1, 30) # time points for time-dependent AUROC (1 to 29 months)
y_train_struct_full = Surv.from_arrays(event=df['death'].to_numpy(dtype=bool), time=df['os_months'])

for fold, (train_idx, test_idx) in enumerate(rskf.split(X, y_event), start=1):
    X_tr, X_te = X.iloc[train_idx], X.iloc[test_idx]
    y_time_tr, y_time_te = y_time.iloc[train_idx], y_time.iloc[test_idx]
    y_event_tr, y_event_te = y_event.iloc[train_idx], y_event.iloc[test_idx]
    
    # Mode imputation
    X_tr_imputed = mode_imputer.fit_transform(X_tr) # fit on training data
    X_te_imputed = mode_imputer.transform(X_te) # transform test data using the same imputer fitted on training data

    # Convert back to DataFrame to preserve column names for CoxPH and other models
    X_tr = pd.DataFrame(X_tr_imputed, columns=X.columns, index=X_tr.index)
    X_te = pd.DataFrame(X_te_imputed, columns=X.columns, index=X_te.index)

    X_tr = pd.get_dummies(X_tr, columns=cat_cols, drop_first=True)
    X_te = pd.get_dummies(X_te, columns=cat_cols, drop_first=True)
    X_tr, X_te = X_tr.align(X_te, join='left', axis=1, fill_value=0)
    X_tr[X_tr.select_dtypes('bool').columns] = X_tr.select_dtypes('bool').astype(int)
    X_te[X_te.select_dtypes('bool').columns] = X_te.select_dtypes('bool').astype(int)
    X_tr = X_tr.apply(pd.to_numeric, errors='coerce')
    X_te = X_te.apply(pd.to_numeric, errors='coerce')

    # structured array for survival models
    y_tr_struct = Surv.from_arrays(event=y_event_tr.to_numpy(dtype=bool), time=y_time_tr)
    y_te_struct = Surv.from_arrays(event=y_event_te.to_numpy(dtype=bool), time=y_time_te)

    # CoxPH for ELN2017 classification with ElasticNet penalty (alpha=0.1, l1_ratio=0.5)
    cox_eln_tr = X_tr[['eln2017mode']].copy()
    cox_eln_tr['os_months'] = y_time_tr
    cox_eln_tr['event'] = y_event_tr
    cph_eln_cv = CoxPHFitter(penalizer=0.1, l1_ratio=0.5) # ElasticNet penalty with alpha=0.1 and l1_ratio=0.5
    cph_eln_cv.fit(cox_eln_tr, duration_col='os_months', event_col='event')
    scores_eln = cph_eln_cv.predict_partial_hazard(X_te[['eln2017mode']])
    cindex_eln_cv.append(concordance_index(y_time_te, -scores_eln)) # high score = lower survival
    fold_eln_cv, _ = cumulative_dynamic_auc(y_tr_struct, y_te_struct, scores_eln, times)
    auroc_eln_cv.append(fold_eln_cv)

    # CoxPH
    cox_tr = X_tr.copy()
    cox_tr['os_months'] = y_time_tr
    cox_tr['event'] = y_event_tr
    cph_cv = CoxPHFitter(penalizer=0.1, l1_ratio=0.5) # ElasticNet penalty with alpha=0.1 and l1_ratio=0.5
    cph_cv.fit(cox_tr, duration_col='os_months', event_col='event')
    scores_cox = cph_cv.predict_partial_hazard(X_te)
    cindex_cox_cv.append(concordance_index(y_time_te, -scores_cox)) # high score = lower survival
    fold_auc_cox, _ = cumulative_dynamic_auc(y_tr_struct, y_te_struct, scores_cox, times)
    auroc_cox_cv.append(fold_auc_cox)
    
    # RSF
    rsf_cv = RandomSurvivalForest(n_estimators=100, min_samples_split=10, min_samples_leaf=15, n_jobs=-1, random_state=10)
    rsf_cv.fit(X_tr, y_tr_struct) # use training data and structured train
    cindex_rsf_cv.append(rsf_cv.score(X_te, y_te_struct)) # c-index on test data
    scores_rsf = rsf_cv.predict(X_te) # get risk scores for test data
    fold_auc_rsf, _ = cumulative_dynamic_auc(y_tr_struct, y_te_struct, scores_rsf, times)
    auroc_rsf_cv.append(fold_auc_rsf)

    # GBM Survival
    gbm_cv = GradientBoostingSurvivalAnalysis(n_estimators=100, learning_rate=0.1, random_state=10)
    gbm_cv.fit(X_tr, y_tr_struct) # use training data and structured train
    cindex_gbm_cv.append(gbm_cv.score(X_te, y_te_struct)) # c-index on test data
    scores_gbm = gbm_cv.predict(X_te) # get risk scores for test data
    fold_auc_gbm, _ = cumulative_dynamic_auc(y_tr_struct, y_te_struct, scores_gbm, times)
    auroc_gbm_cv.append(fold_auc_gbm)

    if fold % 10 == 0:
        print(f"Completed {fold} Cross-validation folds in Survival Analysis...")

# Time-dependent AUROC (AUROC vs. timepoint) for Cox, RSF, GBM survival (full data)
def plot_auroc_over_time(times, listlist_eln, listlist_cox, listlist_rsf, listlist_gbm, put_png=None):
    if put_png is None:
        put_png = os.path.join(FIGURES_DIR, "time_dependent_auroc.jpg")
    # Average the AUCs across all folds (axis=0 means average row-wise)
    mean_eln = np.mean(listlist_eln, axis=0)
    mean_cox = np.mean(listlist_cox, axis=0)
    mean_rsf = np.mean(listlist_rsf, axis=0)
    mean_gbm = np.mean(listlist_gbm, axis=0)

    plt.figure(figsize=(8, 6))
    plt.plot(times, mean_eln, label=f"ELN2017 Cox (Avg AUC: {np.mean(mean_eln):.2f})", lw=2, linestyle='--')
    plt.plot(times, mean_cox, label=f"Cox (Avg AUC: {np.mean(mean_cox):.2f})", lw=2)
    plt.plot(times, mean_rsf, label=f"RSF (Avg AUC: {np.mean(mean_rsf):.2f})", lw=2)
    plt.plot(times, mean_gbm, label=f"GBM (Avg AUC: {np.mean(mean_gbm):.2f})", lw=2)

    plt.xlabel("Months")
    plt.ylabel("Time-dependent AUROC")
    plt.title("Time-dependent AUROC (Cross-Validation Average)")
    plt.ylim(0.4, 1.0)
    plt.legend(loc='lower right')
    plt.grid(True, linestyle='--')
    plt.tight_layout()
    plt.savefig(put_png, dpi=300)
    plt.show()

    return mean_eln, mean_cox, mean_rsf, mean_gbm

print('\nTime-dependent AUROC for Cox, RSF, GBM survival on full data:')
m_eln, m_cox, m_rsf, m_gbm = plot_auroc_over_time(times, auroc_eln_cv, auroc_cox_cv, auroc_rsf_cv, auroc_gbm_cv)

# Summarize Cross-validated results
cindex_cv_data = [
    {
        'Model': 'Cox (ELN2017 classification)',
        'Mean C-index ± SD': f"{np.mean(cindex_eln_cv):.3f} ± {np.std(cindex_eln_cv):.3f}",
        'Avg AUROC': f"{np.mean(m_eln):.3f}"
    },
    {
        'Model': 'Penalized Cox (ElasticNet)', 
        'Mean C-index ± SD': f"{np.mean(cindex_cox_cv):.3f} ± {np.std(cindex_cox_cv):.3f}",
        'Avg AUROC': f"{np.mean(m_cox):.3f}"
    },
    {
        'Model': 'Random Survival Forest', 
        'Mean C-index ± SD': f"{np.mean(cindex_rsf_cv):.3f} ± {np.std(cindex_rsf_cv):.3f}",
        'Avg AUROC': f"{np.mean(m_rsf):.3f}"
    },
    {
        'Model': 'Gradient Boosted Survival (GBM)', 
        'Mean C-index ± SD': f"{np.mean(cindex_gbm_cv):.3f} ± {np.std(cindex_gbm_cv):.3f}",
        'Avg AUROC': f"{np.mean(m_gbm):.3f}"
    }
]

cindex_df = pd.DataFrame(cindex_cv_data)
print("\nC-index for Cox, RSF, GBM survival models (Cross-validated):")
print(cindex_df)

# Save trained survival models
joblib.dump(y_train_struct_full, os.path.join(MODELS_DIR, "y_train_struct.joblib"))
joblib.dump(cph_cv, os.path.join(MODELS_DIR, "cph_model.joblib"))
joblib.dump(rsf_cv, os.path.join(MODELS_DIR, "rsf_model.joblib"))
joblib.dump(gbm_cv, os.path.join(MODELS_DIR, "gbm_model.joblib"))
print("Models saved to:", MODELS_DIR)
# --- Binary models for 12-month survival ---=================================================================
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.neural_network import MLPClassifier
from xgboost import XGBClassifier

# Prepare full dataset for binary outcome
X_binary = df_binary.drop(['12_mo_death'], axis=1)
y_binary = df_binary['12_mo_death']
# drop rows with missing values in the binary outcome (if any)
mask_binary = y_binary.notna()
X_binary = X_binary[mask_binary]
y_binary = y_binary[mask_binary].astype(int) # ensure binary outcome is integer type

# Repeated stratified K-Fold on the full binary dataset
rskf_bin = RepeatedStratifiedKFold(n_splits=5, n_repeats=10, random_state=10)

# Scale numeric features for MLP and logistic regression
from sklearn.preprocessing import StandardScaler
scaler = StandardScaler()
num_cols = ['age_at_diagnosis']

# Cross-validation lists for binary models (LR, RF, XGB, MLP)
aucs_lr, aucs_rf, aucs_xgb, aucs_mlp = [], [], [], []
briers_lr, briers_rf, briers_xgb, briers_mlp = [], [], [], []
probs_lr_all, probs_rf_all, probs_xgb_all, probs_mlp_all = [], [], [], []
y_true_all = []

for fold, (train_idx, test_idx) in enumerate(rskf_bin.split(X_binary, y_binary), start=1):
    Xtr, Xte = X_binary.iloc[train_idx], X_binary.iloc[test_idx]
    ytr, yte = y_binary.iloc[train_idx], y_binary.iloc[test_idx]
    y_true_all.extend(yte) # collect true labels for all folds for later calibration plot
    
    # Mode imputation
    Xtr_imputed = mode_imputer.fit_transform(Xtr) # fit on training data
    Xte_imputed = mode_imputer.transform(Xte) # transform test data using the same imputer fitted on training data
    
    # Convert back to DataFrame to maintain column names for scaling
    Xtr = pd.DataFrame(Xtr_imputed, columns=Xtr.columns, index=Xtr.index)
    Xte = pd.DataFrame(Xte_imputed, columns=Xte.columns, index=Xte.index)
    
    Xtr = pd.get_dummies(Xtr, columns=cat_cols, drop_first=True)
    Xte = pd.get_dummies(Xte, columns=cat_cols, drop_first=True)
    Xtr, Xte = Xtr.align(Xte, join='left', axis=1, fill_value=0)
    Xtr[Xtr.select_dtypes('bool').columns] = Xtr.select_dtypes('bool').astype(int)
    Xte[Xte.select_dtypes('bool').columns] = Xte.select_dtypes('bool').astype(int)
    Xtr = Xtr.apply(pd.to_numeric, errors='coerce')
    Xte = Xte.apply(pd.to_numeric, errors='coerce')

    # scale numeric features for logistic regression and MLP, but not for tree-based models
    Xtr_scaled, Xte_scaled = Xtr.copy(), Xte.copy()
    Xtr_scaled[num_cols] = scaler.fit_transform(Xtr[num_cols])
    Xte_scaled[num_cols] = scaler.transform(Xte[num_cols])

    # Logistic Regression
    lr_cv = LogisticRegression(max_iter=1000, random_state=10)
    lr_cv.fit(Xtr_scaled, ytr) # fit on training scaled data
    probs_lr = lr_cv.predict_proba(Xte_scaled)[:, 1] # get probabilities for the positive class
    aucs_lr.append(auc(*roc_curve(yte, probs_lr)[:2])) # compute AUROC and append
    briers_lr.append(brier_score_loss(yte, probs_lr)) # compute Brier score and append
    probs_lr_all.extend(probs_lr) # collect probabilities for all folds for later calibration plot

    # Random Forest
    rf_cv = RandomForestClassifier(n_estimators=500, random_state=10)
    rf_cv.fit(Xtr, ytr)
    probs_rf = rf_cv.predict_proba(Xte)[:, 1]
    aucs_rf.append(auc(*roc_curve(yte, probs_rf)[:2]))
    briers_rf.append(brier_score_loss(yte, probs_rf))
    probs_rf_all.extend(probs_rf)

    # XGBoost
    xgb_cv = XGBClassifier(n_estimators=200, random_state=10, eval_metric='logloss')
    xgb_cv.fit(Xtr, ytr)
    probs_xgb = xgb_cv.predict_proba(Xte)[:, 1]
    aucs_xgb.append(auc(*roc_curve(yte, probs_xgb)[:2]))
    briers_xgb.append(brier_score_loss(yte, probs_xgb))
    probs_xgb_all.extend(probs_xgb)

    # MLP
    mlp_cv = MLPClassifier(hidden_layer_sizes=(32, 16), max_iter=1000, random_state=10)
    mlp_cv.fit(Xtr_scaled, ytr)
    probs_mlp = mlp_cv.predict_proba(Xte_scaled)[:, 1]
    aucs_mlp.append(auc(*roc_curve(yte, probs_mlp)[:2]))
    briers_mlp.append(brier_score_loss(yte, probs_mlp))
    probs_mlp_all.extend(probs_mlp)

    if fold % 10 == 0:
        print(f"Completed {fold} Cross-validation folds in Binary Classification...")

# ROC curves and AUCs on cross-validation for binary models
def plot_roc_curves_binary(y_true, lr_p, rf_p, xgb_p, mlp_p, put_png=None):
    if put_png is None:
        put_png = os.path.join(FIGURES_DIR, "roc_curves_binary_models.jpg")
    plt.figure(figsize=(8, 6))
    
    # Mapping labels to the probability lists collected in the loop
    plot_data = [
        ('LR', lr_p), ('RF', rf_p), ('XGB', xgb_p), ('MLP', mlp_p)
    ]

    for label, p in plot_data:
        fpr, tpr, _ = roc_curve(y_true, p) # compute FPR and TPR for each model
        plt.plot(fpr, tpr, lw=2, label=f'{label} (AUC = {auc(fpr, tpr):.3f})')

    plt.plot([0, 1], [0, 1], 'k--', alpha=0.5)
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('Cross-validated ROC Curves for Binary Models')
    plt.legend(loc='lower right')
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(put_png, dpi=300)
    plt.show()

print('\nPlotting cross-validated ROC curves for binary models...')
plot_roc_curves_binary(y_true_all, probs_lr_all, probs_rf_all, probs_xgb_all, probs_mlp_all)

# --- AUROC table for binary models using cross-validation results ---
metrics = []
metrics.append({'Model': 'Logistic Regression', 'CV AUROC Mean': round(np.mean(aucs_lr), 3), 'CV AUROC Std': round(np.std(aucs_lr), 3), 'CV Brier Mean': round(np.mean(briers_lr), 3), 'CV Brier Std': round(np.std(briers_lr), 3)})
metrics.append({'Model': 'RandomForest', 'CV AUROC Mean': round(np.mean(aucs_rf), 3), 'CV AUROC Std': round(np.std(aucs_rf), 3), 'CV Brier Mean': round(np.mean(briers_rf), 3), 'CV Brier Std': round(np.std(briers_rf), 3)})
metrics.append({'Model': 'XGBoost', 'CV AUROC Mean': round(np.mean(aucs_xgb), 3), 'CV AUROC Std': round(np.std(aucs_xgb), 3), 'CV Brier Mean': round(np.mean(briers_xgb), 3), 'CV Brier Std': round(np.std(briers_xgb), 3)})
metrics.append({'Model': 'MLP', 'CV AUROC Mean': round(np.mean(aucs_mlp), 3), 'CV AUROC Std': round(np.std(aucs_mlp), 3), 'CV Brier Mean': round(np.mean(briers_mlp), 3), 'CV Brier Std': round(np.std(briers_mlp), 3)})
metrics_df = pd.DataFrame(metrics)

print("\nCross-validated metrics for binary models:")
print(metrics_df.to_string(index=False))

# --- Calibration plots for each binary model: 4 subplots, Lowess smoothing ---
from sklearn.calibration import CalibrationDisplay
fig, axes = plt.subplots(2, 2, figsize=(12, 10)) # 2 rows, 2 columns of subplots
plot_configs = [
    ('Logistic Regression', probs_lr_all),
    ('Random Forest', probs_rf_all),
    ('XGBoost', probs_xgb_all),
    ('MLP', probs_mlp_all)
]

for ax, (name, probs) in zip(axes.flat, plot_configs):
    CalibrationDisplay.from_predictions(
        y_true_all, # true labels collected across all folds
        probs, # predicted probabilities collected across all folds
        n_bins=10,
        ax=ax, 
        name=name
    )
    ax.set_title(f'Calibration: {name}')
    ax.grid(True, linestyle='--', alpha=0.6)

plt.tight_layout()
plt.savefig(os.path.join(FIGURES_DIR, "calibration_1yr_models.jpg"), dpi=300)
plt.show()

# Save trained binary models and scaler
joblib.dump(lr_cv, os.path.join(MODELS_DIR, "logistic_regression.joblib"))
joblib.dump(rf_cv, os.path.join(MODELS_DIR, "random_forest.joblib"))
joblib.dump(xgb_cv, os.path.join(MODELS_DIR, "xgboost_clf.joblib"))
joblib.dump(mlp_cv, os.path.join(MODELS_DIR, "mlp_clf.joblib"))
joblib.dump(scaler, os.path.join(MODELS_DIR, "scaler.joblib"))

print("Binary models saved to:", MODELS_DIR)

# --- Compute model coefficients / importances ---
print('\nComputing model coefficients / importances...')
features = pd.get_dummies(X, columns=cat_cols, drop_first=True).columns

# Extraction using getattr(object, attribute, default)
cox_coefs = cph_cv.params_.reindex(features)
# Convert Cox coefficients to Hazard Ratios (exp(coef)) and compute 95% CI
cox_hr = np.exp(cox_coefs)
if hasattr(cph_cv, 'variance_matrix_'):
    var_mat = cph_cv.variance_matrix_.reindex(index=features, columns=features)
    cox_se = np.sqrt(np.diag(var_mat))
else:
    # fallback: try to read standard errors if available
    cox_se = cph_cv.standard_errors_.reindex(features) if hasattr(cph_cv, 'standard_errors_') else pd.Series([np.nan]*len(features), index=features)

cox_hr_lower = np.exp(cox_coefs - 1.96 * cox_se)
cox_hr_upper = np.exp(cox_coefs + 1.96 * cox_se)

log_coefs = pd.Series(lr_cv.coef_[0], index=features)
# Convert logistic coefficients to Odds Ratios (exp(coef)) and compute approximate 95% CI
log_or = np.exp(log_coefs)

# Approximate standard errors for logistic coefficients using the observed information (Hessian)
try:
    X_design = Xtr_scaled.copy()
    if 'const' not in X_design.columns:
        X_design = X_design.copy()
        X_design.insert(0, 'const', 1.0)
    X_mat = X_design.values
    p = lr_cv.predict_proba(Xtr_scaled)[:, 1]
    W = p * (1 - p)
    XtWX = X_mat.T @ (W[:, None] * X_mat)
    cov_mat = np.linalg.inv(XtWX)
    se_full = np.sqrt(np.diag(cov_mat))
    # drop intercept se for feature-level SEs
    log_se = pd.Series(se_full[1:], index=features)
except Exception:
    log_se = pd.Series([np.nan] * len(features), index=features)

log_or_lower = np.exp(log_coefs - 1.96 * log_se)
log_or_upper = np.exp(log_coefs + 1.96 * log_se)

#rf_imp = pd.Series(rf_cv.feature_importances_, index=features)
xgb_imp = pd.Series(xgb_cv.feature_importances_, index=features)
gbm_imp = pd.Series(gbm_cv.feature_importances_, index=features)

# MLP: Proxy from input-layer absolute mean weights
mlp_weights = np.mean(np.abs(mlp_cv.coefs_[0]), axis=1) if hasattr(mlp_cv, 'coefs_') else [np.nan]*len(features)
mlp_imp = pd.Series(mlp_weights, index=features)

# Consolidated DataFrames
coeffs_df = pd.DataFrame({
    'Variable': features,
    'Cox_HR': cox_hr.values,
    'Cox_HR_CI_lower': cox_hr_lower.values,
    'Cox_HR_CI_upper': cox_hr_upper.values,
    'Logistic_OR': log_or.values,
    'Logistic_OR_CI_lower': log_or_lower.values,
    'Logistic_OR_CI_upper': log_or_upper.values,
}).round(4)

imps_df = pd.DataFrame({
    'Variable': features,
    'XGB': xgb_imp.values,
    'GBM': gbm_imp.values,
    'MLP': mlp_imp.values
}).round(4)

print('\nModel coefficients (Cox Hazard Ratios / Logistic Odds Ratios) for each variable:')
print(coeffs_df)

print('\nModel importance scores for each variable:')
print(imps_df)

# --- Step 13: Robustness Analysis (1,000 Bootstrap Replicates) ---
print("\nStarting 1000 bootstrap replicates...")
n_iterations = 1000
bootstrap_data = []

np.random.seed(10)

for i in range(n_iterations):
    row = {}
    
    # Survival Models Resampling ---
    idx = np.random.choice(len(X), size=len(X), replace=True)
    X_bs = X.iloc[idx].reset_index(drop=True)
    y_t_bs = y_time.iloc[idx].reset_index(drop=True)
    y_e_bs = y_event.iloc[idx].reset_index(drop=True)

    # Mode imputation on the bootstrap sample (in case there are new missing values after resampling)
    X_bs = pd.DataFrame(mode_imputer.fit_transform(X_bs), columns=X.columns)

    #encode
    X_bs = pd.get_dummies(X_bs, columns=cat_cols, drop_first=True)
    X_bs = X_bs.reindex(columns=features, fill_value=0)
    X_bs[X_bs.select_dtypes('bool').columns] = X_bs.select_dtypes('bool').astype(int)
    X_bs = X_bs.apply(pd.to_numeric, errors='coerce')

    # Cox
    cph_bs = CoxPHFitter(penalizer=0.1, l1_ratio=0.5).fit(
        pd.concat([X_bs, y_t_bs, y_e_bs], axis=1), duration_col='os_months', event_col='death'
    )
    cox_coefs = cph_bs.params_.reindex(features)
    cox_hr_bs = np.exp(cox_coefs)
    cindex_cox_bs = cph_bs.concordance_index_

    y_str_bs = Surv.from_arrays(event=y_e_bs.to_numpy(dtype=bool), time=y_t_bs)
    # RSF & GBM Survival
    rsf_bs = RandomSurvivalForest(n_estimators=100, n_jobs=-1).fit(X_bs, y_str_bs)
    cindex_rsf_bs = rsf_bs.score(X_bs, y_str_bs)
    aucs_rsf_bs = rsf_bs.score(X_bs, y_str_bs)
    gbm_bs = GradientBoostingSurvivalAnalysis(n_estimators=100).fit(X_bs, y_str_bs)
    gbm_series = pd.Series(gbm_bs.feature_importances_, index=X_bs.columns)
    cindex_gbm_bs = gbm_bs.score(X_bs, y_str_bs)

    # --- 2. Binary Models Resampling ---
    idxb = np.random.choice(len(X_binary), size=len(X_binary), replace=True)
    X_bi_bs = X_binary.iloc[idxb].reset_index(drop=True)
    y_bi_bs = y_binary.iloc[idxb].reset_index(drop=True)

    # Mode imputation on the bootstrap sample for binary models
    X_bi_bs = pd.DataFrame(mode_imputer.fit_transform(X_bi_bs), columns=X_binary.columns)
    # encode
    X_bi_bs = pd.get_dummies(X_bi_bs, columns=cat_cols, drop_first=True)
    X_bi_bs = X_bi_bs.reindex(columns=features, fill_value=0)
    X_bi_bs[X_bi_bs.select_dtypes('bool').columns] = X_bi_bs.select_dtypes('bool').astype(int)
    X_bi_bs = X_bi_bs.apply(pd.to_numeric, errors='coerce')

    scaler_bs = StandardScaler()
    X_bi_sc = pd.DataFrame(scaler_bs.fit_transform(X_bi_bs), columns=X_bi_bs.columns)

    # Logistic Regression, Random Forest, XGBoost, MLP
    lr_bs = LogisticRegression(max_iter=1000).fit(X_bi_sc, y_bi_bs)
    aucs_lr_bs = auc(*roc_curve(y_bi_bs, lr_bs.predict_proba(X_bi_sc)[:, 1])[:2])
    rf_bs = RandomForestClassifier(n_estimators=500).fit(X_bi_bs, y_bi_bs)
    aucs_rf_bs = auc(*roc_curve(y_bi_bs, rf_bs.predict_proba(X_bi_bs)[:, 1])[:2])
    xgb_bs = XGBClassifier(n_estimators=200, verbosity=0).fit(X_bi_bs, y_bi_bs)
    aucs_xgb_bs = auc(*roc_curve(y_bi_bs, xgb_bs.predict_proba(X_bi_bs)[:, 1])[:2])
#    mlp_bs = MLPClassifier(hidden_layer_sizes=(32, 16), max_iter=1000).fit(X_bi_sc, y_bi_bs)

    # --- 3. Collect Results ---
    log_series = pd.Series(lr_bs.coef_.ravel(), index=X_bi_sc.columns)
    log_or_bs = np.exp(log_series)
    rf_series = pd.Series(rf_bs.feature_importances_, index=X_bi_bs.columns)
    xgb_series = pd.Series(xgb_bs.feature_importances_, index=X_bi_bs.columns)
#    mlp_series = pd.Series(np.mean(np.abs(mlp_bs.coefs_[0]), axis=1), index=X_bi_sc.columns)

    for f in features:
        row[f'Cox_{f}']      = cox_hr_bs.get(f)
        row[f'Logistic_{f}'] = log_or_bs.get(f)
        row[f'RF_{f}']       = rf_series.get(f)
        row[f'XGB_{f}']      = xgb_series.get(f)
        row[f'GBM_{f}']      = gbm_series.get(f)
#        row[f'MLP_{f}']      = mlp_series.get(f)

    # Overall model performance metrics for bootstrap
    row['Cindex_Cox'] = cindex_cox_bs
    row['Cindex_RSF'] = cindex_rsf_bs
    row['Cindex_GBM'] = cindex_gbm_bs
    row['AUROC_Logistic'] = aucs_lr_bs
    row['AUROC_RF'] = aucs_rf_bs
    row['AUROC_XGB'] = aucs_xgb_bs

    bootstrap_data.append(row)
    if (i + 1) % 100 == 0: print(f"Progress: {i+1}/{n_iterations}")

# --- Summarize Bootstrap Results ---
bs_df = pd.DataFrame(bootstrap_data)

# General robust summary across all variables and models (mean and 95% CI)
summary = bs_df.agg(['mean', lambda x: x.quantile(0.025), lambda x: x.quantile(0.975)]).T
summary.columns = ['Mean', '95% CI Lower', '95% CI Upper']
robustness_summary_df = summary.round(3).reset_index().rename(columns={'index': 'Metric'})

# Overall model performance metrics from bootstrap
cindex_cols = ['Cindex_Cox', 'Cindex_RSF', 'Cindex_GBM']
auroc_cols = ['AUROC_Logistic', 'AUROC_RF', 'AUROC_XGB']

cindex_bs = bs_df[cindex_cols]
cindex_summary = cindex_bs.agg(['mean', lambda x: x.quantile(0.025), lambda x: x.quantile(0.975)]).T
cindex_summary.columns = ['Mean', '95% CI Lower', '95% CI Upper']
cindex_summary_df = cindex_summary.round(3).reset_index().rename(columns={'index': 'Metric'})

auroc_bs = bs_df[auroc_cols]
auroc_summary = auroc_bs.agg(['mean', lambda x: x.quantile(0.025), lambda x: x.quantile(0.975)]).T
auroc_summary.columns = ['Mean', '95% CI Lower', '95% CI Upper']
auroc_summary_df = auroc_summary.round(3).reset_index().rename(columns={'index': 'Metric'})

print("\nRobustness Summary (95% Confidence Intervals):")
print(robustness_summary_df.head())
print("\nBootstrap C-index Summary:")
print(cindex_summary_df)
print("\nBootstrap AUROC Summary:")
print(auroc_summary_df)

# --- Step 14: Export tables and figures to .docx (final) ---====================================================
from docx import Document
from docx.shared import Inches
import os

doc = Document()
doc.add_heading('ML Project Results', 0)

# Summary tables
def df_to_doc_table(document, df, title):
    document.add_heading(title, level=1)
    tbl = document.add_table(rows=1, cols=len(df.columns))
    tbl.style = 'Table Grid'
    for i, col in enumerate(df.columns):
        tbl.cell(0, i).text = str(col)
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])

doc.add_heading("Patient Exclusion Summary", level=1)
doc.add_paragraph(f"Total patients: {n_total}")
doc.add_paragraph(f"Excluded due to missing death information: {n_missing_death_info}")
doc.add_paragraph(f"Included in survival analysis: {len(df)}")
doc.add_paragraph(f"Excluded from binary model (censored <12 months, outcome unknown): {uncertain_mask.sum()}")
doc.add_paragraph(f"Included in binary models: {len(df_binary)}")
df_to_doc_table(doc, summary_df, 'Summary Table: Age at Diagnosis and OS Months')
df_to_doc_table(doc, gene_summary_df, 'Gene Mutation Summary Table')
df_to_doc_table(doc, sex_summary, 'Sex Summary Table')
df_to_doc_table(doc, denovo_summary, 'Denovo Category Summary Table')
df_to_doc_table(doc, eln_summary, 'ELN 2017 Classification Summary Table')
df_to_doc_table(doc, chrom_summary, 'Chromosome Category Summary Table')
df_to_doc_table(doc, cindex_df, 'Survival Models Harrell\'s C-index Comparison and AUROC at 12 months')
df_to_doc_table(doc, metrics_df, 'Binary Models Performance (AUROCs and Brier scores)')
df_to_doc_table(doc, coeffs_df, 'Model Hazard Ratios for each variables')
df_to_doc_table(doc, imps_df, 'Model Importance Scores for each variable')
df_to_doc_table(doc, robustness_summary_df, 'Robustness Analysis: 1000 Bootstrap 95% CIs')
df_to_doc_table(doc, cindex_summary_df, 'Bootstrap Summary: C-index (Cox / RSF / GBM)')
df_to_doc_table(doc, auroc_summary_df, 'Bootstrap Summary: AUROC (Logistic / RF / XGB)')

# Figures to add
def add_fig(document, path, caption):
    document.add_heading(caption, level=1)
    if os.path.exists(path):
        document.add_picture(path, width=Inches(5.5))
    else:
        document.add_paragraph(f'Figure not found: {path}')

    add_fig(doc, os.path.join(FIGURES_DIR, 'time_dependent_auroc.jpg'),
        'Time-dependent AUROC for Cox, RSF, GBM survival')
    add_fig(doc, os.path.join(FIGURES_DIR, 'roc_curves_binary_models.jpg'),
        'ROC Curves for Binary Models (Logistic Regression, Random Forest, XGBoost, MLP)')
    add_fig(doc, os.path.join(FIGURES_DIR, 'calibration_1yr_models.jpg'),
        'Calibration Plots for Binary Models (Logistic Regression, Random Forest, XGBoost, MLP)')

output_path = os.path.join(OUTPUT_DIR, "ML_all_results.docx") # save the document in the output directory
doc.save(output_path) # save the document
print(f"\nAll results have been exported to {output_path}")
plt.close('all')
