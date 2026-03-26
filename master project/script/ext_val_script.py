# Master Project - External Validation Script
# Machine Learning for Survival Analysis in AML
# Python version 3.12
# 03/23/2026
# by Steve Tungjitviboonkun

# Required packages are listed in requirements.txt
# Install them using: pip install -r requirements.txt

#Master Project/
#├── data/
#│   └── data_clinical.csv
#│   └── data_clinical_template.csv
#├── script/
#│   └── main_script.py
#│   └── ext_val_script.py
#├── output/
#│   ├── models/
#│   └── figures/

# Pipeline:
# 1. Load and clean data
# 2. Survival modeling with Cox, RSF, GBM
# 3. Binary modeling for 12-month mortality
# 4. Bootstrap robustness analysis
# 5. Export results to DOCX

import os
import re
import pandas as pd
import numpy as np
import joblib
from datetime import datetime
from sklearn.metrics import roc_curve, auc, brier_score_loss
from sklearn.calibration import CalibrationDisplay
from lifelines.utils import concordance_index
from sksurv.util import Surv
from sksurv.metrics import cumulative_dynamic_auc
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from lifelines import CoxPHFitter
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from xgboost import XGBClassifier

try:
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
except NameError:
    BASE_DIR = os.getcwd()

# Paths
DATA_PATH = os.path.join(BASE_DIR, "data", "data_clinical_template.csv")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
MODELS_DIR = os.path.join(OUTPUT_DIR, "models")
FIGURES_DIR = os.path.join(OUTPUT_DIR, "figures_ext_val")
OUT_DOC = os.path.join(OUTPUT_DIR, "ML_ext_val_results.docx")

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(MODELS_DIR, exist_ok=True)
os.makedirs(FIGURES_DIR, exist_ok=True)

# Step 1: Load data
df = pd.read_csv(DATA_PATH)
print(df.head())

# Step 2: Data Cleaning - Parse dates (handles month/day/year numeric formats)
df['date_last'] = pd.to_datetime(df['date_last'], errors='coerce')
df['date_dx'] = pd.to_datetime(df['date_dx'], errors='coerce')

# Create os_months as months between date_last and date_dx
df['os_months'] = (df['date_last'] - df['date_dx']).dt.days / 30.44

n_total = len(df)
print('Total patients in external dataset:', n_total)
# missing death information (cannot be included in survival analysis)
n_missing_death_info = df['death'].isna().sum()
print('Patients with missing death information:', n_missing_death_info)
df = df[df['death'].notna()]

# drop if os_months < 0 since this is likely an error and we can't use them for survival analysis
n_negative_os = (df['os_months'] <= 0).sum()
n_missing_os = df['os_months'].isna().sum()
print('Patients with negative OS months (invalid):', n_negative_os)
print('Patients with missing OS months:', n_missing_os)
df = df[df['os_months'] > 0]

print('Patients included in survival analysis:', len(df))
# Exclude from binary model if censored <12 months (outcome unknown)
uncertain_mask = ((df['death'] == 'no') | (df['death'] == 'No') | (df['death'] == 'NO')) & (df['os_months'] < 12)
print('Patients excluded from binary model due to uncertain 12-month outcome:', uncertain_mask.sum())

# state the predictors, outcomes we will use
predictors = [
    'age_at_diagnosis', 'sex', 'eln2017mode', 'denovo_cat', 'chromosome_cat',
    'flt3_itd_cat', 'npm1_cat', 'runx1_cat', 'asxl1_cat', 'tp53_cat'
    ]
outcomes = ['os_months', 'death']

# set every variable to lowercase and strip whitespace to match the main dataset formatting
for col in df.columns:
    if df[col].dtype == 'object' or df[col].dtype == 'string':
        df[col] = df[col].astype(str).str.lower().str.strip()

# change numeric to numeric type, coerce errors to NaN (for age_at_dx and os_months)
df['age_at_dx'] = pd.to_numeric(df['age_at_dx'], errors='coerce')
df['age_at_diagnosis'] = pd.to_numeric(df['age_at_dx'], errors='coerce')
df['os_months'] = pd.to_numeric(df['os_months'], errors='coerce')

# Match eln2017mode with eln_class
df['eln2017mode'] = df['eln_class'].copy()

# Map common combined tokens to consistent labels
df.loc[df['eln2017mode'] == 'favorableorintermediate', 'eln2017mode'] = 'intermediate'
df.loc[df['eln2017mode'] == 'intermediateoradverse', 'eln2017mode'] = 'adverse'
df.loc[df['eln2017mode'] == 'missingkaryo', 'eln2017mode'] = np.nan
df.loc[df['eln2017mode'] == 'missingmutations', 'eln2017mode'] = np.nan
df.loc[df['eln2017mode'] == 'noninitial', 'eln2017mode'] = np.nan

# Match denovo_cat with is_denovo
df['denovo_cat'] = df['is_denovo'].copy()
df.loc[df['denovo_cat'] == 'missing', 'denovo_cat'] = np.nan

# Clean karyotype -> chromosome_cat (0 normal for 46,XX[20] or 46,XY[20], else 1 complex)
def classify_karyo(val):
    if pd.isna(val):
        return np.nan
    s = str(val).lower().replace(' ', '')
    # match patterns like 46,xx[20] or 46,xy[20]
    if re.match(r'^46,xx\[20\]$', s) or re.match(r'^46,xy\[20\]$', s):
        return 'normal'
    return 'complex'

df['chromosome_cat'] = df['karyotype'].apply(classify_karyo)

# Keep only predictors or outcomes columns
df = df.loc[:, predictors + outcomes]
print(df.head())

# Step 3: Summary Statistics for age_at_diagnosis and os_months
summary_vars = ['age_at_diagnosis', 'os_months']
summary_table = []
for var in summary_vars:
    col = df[var]
    mean = round(col.mean(), 2)
    median = round(col.median(), 2)
    min_ = round(col.min(), 2)
    max_ = round(col.max(), 2)
    n_missing = col.isna().sum()
    pct_missing = round(100 * n_missing / len(col), 2)
    summary_table.append({
        'Variable': var,
        'Mean': mean,
        'Median': median,
        'Min': min_,
        'Max': max_,
        'Missing (n)': n_missing,
        'Missing (%)': pct_missing
    })
summary_df = pd.DataFrame(summary_table)
print("\nSummary table for age_at_diagnosis and os_months:")
print(summary_df)

# Gene mutation summary
gene_cols = ['flt3_itd_cat', 'npm1_cat', 'runx1_cat', 'asxl1_cat', 'tp53_cat']
gene_summary = []
for col in gene_cols:
    value_counts = df[col].value_counts(dropna=False)
    pos = value_counts.get('positive', 0)
    neg = value_counts.get('negative', 0)
    missing = value_counts.get(np.nan, 0) if np.nan in value_counts else df[col].isna().sum()
    total = len(df)
    gene_summary.append({
        'Gene': col,
        'Positive (n)': pos,
        'Positive (%)': round(100 * pos / total, 2),
        'Negative (n)': neg,
        'Negative (%)': round(100 * neg / total, 2),
        'Missing (n)': missing,
        'Missing (%)': round(100 * missing / total, 2)
    })
gene_summary_df = pd.DataFrame(gene_summary)
print("\nGene mutation summary table:")
print(gene_summary_df)

# Sex summary
sex_counts = df['sex'].value_counts(dropna=False)
total = len(df)
male = sex_counts.get('male', 0)
female = sex_counts.get('female', 0)
missing = df['sex'].isna().sum()
total = len(df)
sex_summary = pd.DataFrame({
    'Sex': ['male', 'female', 'missing', 'total'],
    'Count': [male, female, missing, total],
    'Percent': [round(100 * male / total, 2), round(100 * female / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\nSex summary table:")
print(sex_summary)

# Table for denovo_cat (True/False/Missing)
denovo_counts = df['denovo_cat'].value_counts(dropna=False)
total = len(df)
true_ = denovo_counts.get('yes', 0)
false_ = denovo_counts.get('no', 0)
missing = df['denovo_cat'].isna().sum()
denovo_summary = pd.DataFrame({
    'denovo_cat': ['yes', 'no', 'missing', 'total'],
    'Count': [true_, false_, missing, total],
    'Percent': [round(100 * true_ / total, 2), round(100 * false_ / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\ndenovo_cat summary table:")
print(denovo_summary)

# Table for eln2017mode (favorable/intermediate/adverse/missing)
eln_counts = df['eln2017mode'].value_counts(dropna=False)
total = len(df)
fav = eln_counts.get('favorable', 0)
intm = eln_counts.get('intermediate', 0)
advs = eln_counts.get('adverse', 0)
missing = df['eln2017mode'].isna().sum()
eln_summary = pd.DataFrame({
    'ELN class': ['favorable', 'intermediate', 'adverse', 'missing', 'total'],
    'Count': [fav, intm, advs, missing, total],
    'Percent': [round(100 * fav / total, 2), round(100 * intm / total, 2), round(100 * advs / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\nELN 2017 classification summary table:")
print(eln_summary)

# check mean os_months by eln2017mode
os_by_eln = df.groupby('eln2017mode')['os_months'].agg(['mean', 'median', 'min', 'max', 'count']).reset_index().round(2)
print("\nOS months by ELN2017 mode:")
print(os_by_eln)

# Table for chromosome_cat (normal/complex/missing)
chrom_counts = df['chromosome_cat'].value_counts(dropna=False)
total = len(df)
normal = chrom_counts.get('normal', 0)
complex_ = chrom_counts.get('complex', 0)
missing = df['chromosome_cat'].isna().sum()
chrom_summary = pd.DataFrame({
    'chromosome cat': ['normal', 'complex', 'missing', 'total'],
    'Count': [normal, complex_, missing, total],
    'Percent': [round(100 * normal / total, 2), round(100 * complex_ / total, 2), round(100 * missing / total, 2), 100.0]
})
print("\nchromosome category summary table:")
print(chrom_summary)

# Prepare data for model evaluation
cat_cols = [
    'sex', 'chromosome_cat', 'denovo_cat',
    'flt3_itd_cat', 'npm1_cat', 'runx1_cat',
    'asxl1_cat', 'tp53_cat'
]

df['eln2017mode'] = df['eln2017mode'].map({'favorable': 0, 'intermediate': 1, 'adverse': 2})
df['death'] = df['death'].map({'no': 0, 'yes': 1})
print(df.dtypes)
# check mean os_months by eln2017mode
df.groupby('eln2017mode')['os_months'].agg(['mean', 'median', 'min', 'max', 'count']).reset_index()

# Create binary dataframe for 12-month mortality
df_binary = df[~uncertain_mask].copy()
df_binary['12_mo_death']=(df_binary['os_months'] < 12).astype(int)
df_binary = df_binary[predictors + ['12_mo_death']] # binary model
print(df_binary.head())
print(df_binary.dtypes)
print('Patients included in binary analysis:', len(df_binary))

# Load all models ===================================================================================================
models = {}
model_files = {
    'cph': os.path.join(MODELS_DIR, 'cph_model.joblib'),
    'rsf': os.path.join(MODELS_DIR, 'rsf_model.joblib'),
    'gbm': os.path.join(MODELS_DIR, 'gbm_model.joblib'),
    'lr': os.path.join(MODELS_DIR, 'logistic_regression.joblib'),
    'rf': os.path.join(MODELS_DIR, 'random_forest.joblib'),
    'xgb': os.path.join(MODELS_DIR, 'xgboost_clf.joblib'),
    'mlp': os.path.join(MODELS_DIR, 'mlp_clf.joblib'),
    'scaler': os.path.join(MODELS_DIR, 'scaler.joblib')
}
for name, path in model_files.items():
    if os.path.exists(path):
        models[name] = joblib.load(path)
    else:
        print(f'Warning: model file not found: {path}')

# Step 4: Evaluate survival models (C-index and time-dependent AUC) ================================
from sklearn.impute import SimpleImputer
mode_imputer = SimpleImputer(strategy='most_frequent')

X_ext = df[predictors].copy()
y_time = df['os_months']
y_event = df['death']

# Imputation for survival analysis (use mode imputer)
X_ext_imputed = mode_imputer.fit_transform(X_ext)
X_ext = pd.DataFrame(X_ext_imputed, columns=X_ext.columns, index=X_ext.index)

# Create dummy variables for categorical predictors and align to training features
X_ext = pd.get_dummies(X_ext, columns=cat_cols, drop_first=True)
X_ext[X_ext.select_dtypes('bool').columns] = X_ext.select_dtypes('bool').astype(int)
X_ext = X_ext.apply(pd.to_numeric, errors='coerce')

# Extract training feature names from Cox model
if 'cph' in models:
    trained_features = models['cph'].params_.index
else:
    raise ValueError("Cox model required for feature alignment")

X_ext = X_ext.reindex(columns=trained_features, fill_value=0)

# Evaluate survival models using C-index
survival_results = {}
aucs = {}

times = np.arange(3, 30)
y_train_struct = joblib.load(os.path.join(MODELS_DIR, 'y_train_struct.joblib')) # load training from main script
y_struct = Surv.from_arrays(event=y_event.to_numpy(dtype=bool), time=y_time)

# Traditional ELN2017 evaluation (using eln2017mode as risk score)
X_eln = X_ext[['eln2017mode']].copy()
cox_eln_df = X_eln.copy()
cox_eln_df['os_months'] = y_time
cox_eln_df['death'] = y_event
cph_eln = CoxPHFitter()
cph_eln.fit(cox_eln_df, duration_col='os_months', event_col='death')
scores_eln = cph_eln.predict_partial_hazard(X_eln)
survival_results['eln_cox_cindex'] = concordance_index(y_time, -scores_eln)
auc_eln, _ = cumulative_dynamic_auc(y_train_struct, y_struct, scores_eln, times)

aucs['ELN Cox'] = auc_eln

if 'cph' in models:
    try:
        scores_cox = models['cph'].predict_partial_hazard(X_ext)
        survival_results['cox_cindex'] = concordance_index(y_time, -scores_cox)
        auc_cph, _ = cumulative_dynamic_auc(y_train_struct, y_struct, scores_cox, times)
    except Exception as e:
        print('Cox evaluation error:', e)
if 'rsf' in models:
    try:
        y_struct = Surv.from_arrays(event=y_event.to_numpy(dtype=bool), time=y_time)
        survival_results['rsf_cindex'] = models['rsf'].score(X_ext, y_struct)
        scores_rsf = models['rsf'].predict(X_ext)
        auc_rsf, _ = cumulative_dynamic_auc(y_train_struct, y_struct, scores_rsf, times)
    except Exception as e:
        print('RSF evaluation error:', e)
if 'gbm' in models:
    try:
        y_struct = Surv.from_arrays(event=y_event.to_numpy(dtype=bool), time=y_time)
        survival_results['gbm_cindex'] = models['gbm'].score(X_ext, y_struct)
        scores_gbm = models['gbm'].predict(X_ext)
        auc_gbm, _ = cumulative_dynamic_auc(y_train_struct, y_struct, scores_gbm, times)
    except Exception as e:
        print('GBM survival evaluation error:', e)


# Model evaluation results: C-index table and time-dependent AUROC plot
surv_cindex = []
surv_cindex.append({'Model': 'Penalized Cox (ElasticNet)', 'C-index': round(survival_results['cox_cindex'], 3)})
surv_cindex.append({'Model': 'Random Survival Forest', 'C-index': round(survival_results['rsf_cindex'], 3)})
surv_cindex.append({'Model': 'Gradient Boosted Survival Models', 'C-index': round(survival_results['gbm_cindex'], 3)})
cindex_df = pd.DataFrame(surv_cindex)

# Time-dependent AUROC over months (use external data as both train/test)
plt.figure(figsize=(8, 6))
for name, vals in aucs.items():
    plt.plot(times, auc_eln, label=f"ELN2017 Cox (AUROC: {np.mean(auc_eln):.2f})", lw=2, linestyle='--')
    plt.plot(times, auc_cph, label=f"Cox (AUROC: {np.mean(auc_cph):.2f})", lw=2)
    plt.plot(times, auc_rsf, label=f"RSF (AUROC: {np.mean(auc_rsf):.2f})", lw=2)
    plt.plot(times, auc_gbm, label=f"GBM (AUROC: {np.mean(auc_gbm):.2f})", lw=2)

    plt.xlabel('Months')
    plt.ylabel('Time-dependent AUROC')
    plt.title('Time-dependent AUROC (External Validation)')
    plt.ylim(0.4, 1.0)
    plt.legend(loc='lower right')
    plt.grid(True, linestyle='--')
    plt.tight_layout()
    surv_roc_path = os.path.join(FIGURES_DIR, 'time_survival_roc_12m_ext.jpg')
    plt.savefig(surv_roc_path, dpi=300)
    plt.close()

print("\nSurvival model performance metrics (C-index):")
print(cindex_df)

# Step 5: Binary 12-month evaluation =================================================================
aucs_lr, aucs_rf, aucs_xgb, aucs_mlp = [], [], [], []
briers_lr, briers_rf, briers_xgb, briers_mlp = [], [], [], []
probs_lr, probs_rf, probs_xgb, probs_mlp = [], [], [], []
y_true_all = []

Xb = df_binary[predictors].copy()
yb = df_binary['12_mo_death']
y_true_all.extend(yb.tolist()) # collect true labels across all folds for calibration plot

# Imputation for binary analysis (use most frequent / mode imputer)
Xb_imputed = mode_imputer.fit_transform(Xb)
Xb = pd.DataFrame(Xb_imputed, columns=Xb.columns, index=Xb.index)

# One-hot encoding (same as training)
Xb = pd.get_dummies(Xb, columns=cat_cols, drop_first=True)
Xb[Xb.select_dtypes('bool').columns] = Xb.select_dtypes('bool').astype(int)
Xb = Xb.apply(pd.to_numeric, errors='coerce')

# --- align with training features ---
if 'lr' in models:
    trained_features_bin = models['lr'].feature_names_in_
else:
    trained_features_bin = Xb.columns

Xb = Xb.reindex(columns=trained_features_bin, fill_value=0)

# Scale numeric features if scaler exists
if 'scaler' in models:
    scaler = models['scaler']
    Xb_scaled = Xb.copy()
    if 'age_at_diagnosis' in Xb_scaled.columns:
        Xb_scaled[['age_at_diagnosis']] = scaler.transform(Xb_scaled[['age_at_diagnosis']])
else:
    Xb_scaled = Xb.copy()

# --- Model evaluation ---
if 'lr' in models:
    probs_lr = models['lr'].predict_proba(Xb_scaled)[:, 1]
    aucs_lr.append(auc(*roc_curve(yb, probs_lr)[:2]))
    briers_lr.append(brier_score_loss(yb, probs_lr))

if 'rf' in models:
    probs_rf = models['rf'].predict_proba(Xb)[:, 1]
    aucs_rf.append(auc(*roc_curve(yb, probs_rf)[:2]))
    briers_rf.append(brier_score_loss(yb, probs_rf))

if 'xgb' in models:
    probs_xgb = models['xgb'].predict_proba(Xb)[:, 1]
    aucs_xgb.append(auc(*roc_curve(yb, probs_xgb)[:2]))
    briers_xgb.append(brier_score_loss(yb, probs_xgb))

if 'mlp' in models:
    probs_mlp = models['mlp'].predict_proba(Xb_scaled)[:, 1]
    aucs_mlp.append(auc(*roc_curve(yb, probs_mlp)[:2]))
    briers_mlp.append(brier_score_loss(yb, probs_mlp))

# --- AUROC table for binary models using external validation results ---
metrics = []
metrics.append({'Model': 'Logistic Regression', 'AUROC': aucs_lr[0], 'Brier': briers_lr[0]})
metrics.append({'Model': 'RandomForest', 'AUROC': aucs_rf[0], 'Brier': briers_rf[0]})
metrics.append({'Model': 'XGBoost', 'AUROC': aucs_xgb[0], 'Brier': briers_xgb[0]})
metrics.append({'Model': 'MLP', 'AUROC': aucs_mlp[0], 'Brier': briers_mlp[0]})
metrics_df = pd.DataFrame(metrics).round(3)

# Binary models: ROC curve plot
def plot_roc_curves_binary(y_true, lr_p, rf_p, xgb_p, mlp_p, put_png=None):
    if put_png is None:
        put_png = os.path.join(FIGURES_DIR, "roc_curves_binary_models_ext.jpg")
    plt.figure(figsize=(8, 6))
    
    # Mapping labels to the probability lists collected in the loop
    plot_data = [
        ('LR', lr_p), ('RF', rf_p), ('XGB', xgb_p), ('MLP', mlp_p)
    ]

    for label, p in plot_data:
        fpr, tpr, _ = roc_curve(y_true, p) # compute FPR and TPR for each model
        plt.plot(fpr, tpr, lw=2, label=f'{label} (AUC = {auc(fpr, tpr):.3f})')

    plt.plot([0, 1], [0, 1], 'k--', alpha=0.5)
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curves for Binary Models from External Validation')
    plt.legend(loc='lower right')
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(put_png, dpi=300)
    plt.show()

print('\nPlotting ROC curves for binary models using external validation results...')
plot_roc_curves_binary(y_true_all, probs_lr, probs_rf, probs_xgb, probs_mlp)

# Binary models: Calibration plots
fig, axes = plt.subplots(2, 2, figsize=(12, 10)) # 2 rows, 2 columns of subplots
plot_configs = [
    ('Logistic Regression', probs_lr),
    ('Random Forest', probs_rf),
    ('XGBoost', probs_xgb),
    ('MLP', probs_mlp)
]

for ax, (name, probs) in zip(axes.flat, plot_configs):
    CalibrationDisplay.from_predictions(
        y_true_all, # true labels collected across all folds
        probs, # predicted probabilities collected across all folds
        n_bins=10,
        ax=ax, 
        name=name
    )
    ax.set_title(f'Calibration: {name}')
    ax.grid(True, linestyle='--', alpha=0.6)

plt.tight_layout()
plt.savefig(os.path.join(FIGURES_DIR, "calibration_1yr_models_ext.jpg"), dpi=300)
plt.show()

# --- Step 6: Bootstrap ---
print('\nStarting 1000 bootstrap analysis for survival and binary models...')
n_iterations = 1000
np.random.seed(10)

bs_metrics = []

for i in range(n_iterations):
    row = {}

    # sample with replacement for survival evaluation
    idx = np.random.choice(len(X_ext), len(X_ext), replace=True)
    Xb_ext = X_ext.iloc[idx].reset_index(drop=True)
    y_t_bs = y_time.iloc[idx].reset_index(drop=True)
    y_e_bs = y_event.iloc[idx].reset_index(drop=True)
    y_struct_bs = Surv.from_arrays(event=y_e_bs.astype(bool), time=y_t_bs)

    # compute C-index for stored survival models on bootstrap sample
    if 'cph' in models:
        try:
            scores = models['cph'].predict_partial_hazard(Xb_ext)
            row['Cindex_Cox'] = concordance_index(y_t_bs, -scores)
        except:
            row['Cindex_Cox'] = np.nan
    else:
        row['Cindex_Cox'] = np.nan

    if 'rsf' in models:
        try:
            row['Cindex_RSF'] = models['rsf'].score(Xb_ext, y_struct_bs)
        except:
            row['Cindex_RSF'] = np.nan
    else:
        row['Cindex_RSF'] = np.nan

    if 'gbm' in models:
        try:
            row['Cindex_GBM'] = models['gbm'].score(Xb_ext, y_struct_bs)
        except:
            row['Cindex_GBM'] = np.nan
    else:
        row['Cindex_GBM'] = np.nan

    # binary bootstrap sample
    if len(df_binary) > 0:
        idxb = np.random.choice(len(df_binary), len(df_binary), replace=True)
        df_bi_bs = df_binary.iloc[idxb].reset_index(drop=True)
        yb_bs = df_bi_bs['12_mo_death']

        Xb_bs = pd.get_dummies(pd.DataFrame(mode_imputer.transform(df_bi_bs[predictors]), columns=predictors), columns=cat_cols, drop_first=True)
        Xb_bs = Xb_bs.reindex(columns=trained_features_bin, fill_value=0)
        Xb_bs[Xb_bs.select_dtypes('bool').columns] = Xb_bs.select_dtypes('bool').astype(int)
        Xb_bs = Xb_bs.apply(pd.to_numeric, errors='coerce')

        Xb_scaled_bs = Xb_bs.copy()
        if 'scaler' in models and 'age_at_diagnosis' in Xb_scaled_bs.columns:
            Xb_scaled_bs[['age_at_diagnosis']] = models['scaler'].transform(Xb_scaled_bs[['age_at_diagnosis']])

        # compute AUROC for stored binary models on bootstrap sample
        try:
            row['AUROC_Logistic'] = auc(*roc_curve(yb_bs, models['lr'].predict_proba(Xb_scaled_bs)[:, 1])[:2]) if 'lr' in models else np.nan
        except:
            row['AUROC_Logistic'] = np.nan

        try:
            row['AUROC_RF'] = auc(*roc_curve(yb_bs, models['rf'].predict_proba(Xb_bs)[:, 1])[:2]) if 'rf' in models else np.nan
        except:
            row['AUROC_RF'] = np.nan

        try:
            row['AUROC_XGB'] = auc(*roc_curve(yb_bs, models['xgb'].predict_proba(Xb_bs)[:, 1])[:2]) if 'xgb' in models else np.nan
        except:
            row['AUROC_XGB'] = np.nan
    else:
        row['AUROC_Logistic'] = row['AUROC_RF'] = row['AUROC_XGB'] = np.nan

    bs_metrics.append(row)

    if (i + 1) % 100 == 0:
        print(f'Bootstrap progress: {i+1}/{n_iterations}')

bs_metrics_df = pd.DataFrame(bs_metrics)

def summarize_metrics(df, cols):
    out = df[cols].agg(['mean', lambda x: x.quantile(0.025), lambda x: x.quantile(0.975)]).T
    out.columns = ['Mean', '95% CI Lower', '95% CI Upper']
    return out.round(3).reset_index().rename(columns={'index': 'Metric'})

cindex_cols = ['Cindex_Cox', 'Cindex_RSF', 'Cindex_GBM']
cindex_summary_df = summarize_metrics(bs_metrics_df, cindex_cols)

print('\nBootstrap Summary: C-index (Cox / RSF / GBM)')
print(cindex_summary_df)

auroc_cols = ['AUROC_Logistic', 'AUROC_RF', 'AUROC_XGB']
auroc_summary_df = summarize_metrics(bs_metrics_df, auroc_cols)

print('\nBootstrap Summary: AUROC (Logistic / RF / XGB)')
print(auroc_summary_df)

# Step 7:Final document assembly and save ==============================
doc = Document()
doc.add_heading('External Validation Results', 0)

# Summary Table
def df_to_doc_table(document, df, title):
    document.add_heading(title, level=1)
    tbl = document.add_table(rows=1, cols=len(df.columns))
    tbl.style = 'Table Grid'
    for i, col in enumerate(df.columns):
        tbl.cell(0, i).text = str(col)
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])

doc.add_heading("Patient Exclusion Summary", level=1)
doc.add_paragraph(f"Total patients: {n_total}")
doc.add_paragraph(f"Excluded due to missing death information: {n_missing_death_info}")
doc.add_paragraph(f"Excluded due to negative OS months (invalid): {n_negative_os}")
doc.add_paragraph(f"Excluded due to missing OS months: {n_missing_os}")
doc.add_paragraph(f"Included in survival analysis: {len(df)}")
doc.add_paragraph(f"Excluded from binary model (censored <12 months, outcome unknown): {uncertain_mask.sum()}")
doc.add_paragraph(f"Included in binary models: {len(df_binary)}")
df_to_doc_table(doc, summary_df, 'Summary Table: Age at Diagnosis and OS Months')
df_to_doc_table(doc, gene_summary_df, 'Gene Mutation Summary Table')
df_to_doc_table(doc, sex_summary, 'Sex Summary Table')
df_to_doc_table(doc, denovo_summary, 'Denovo Category Summary Table')
df_to_doc_table(doc, eln_summary, 'ELN 2017 Classification Summary Table')
df_to_doc_table(doc, os_by_eln, 'OS Months by ELN2017 Classification Table')
df_to_doc_table(doc, chrom_summary, 'Chromosome Category Summary Table')
df_to_doc_table(doc, cindex_df, 'Survival Models Harrell\'s C-index Comparison and AUROC at 12 months')
df_to_doc_table(doc, metrics_df, 'Binary Models Performance (AUROCs and Brier scores)')
df_to_doc_table(doc, cindex_summary_df, '1000 Bootstrap Summary: C-index (Cox / RSF / GBM)')
df_to_doc_table(doc, auroc_summary_df, '1000 Bootstrap Summary: AUROC (Logistic / RF / XGB)')

# Figures to add
def add_fig(document, path, caption):
    document.add_heading(caption, level=1)
    if os.path.exists(path):
        document.add_picture(path, width=Inches(5.5))
    else:
        document.add_paragraph(f'Figure not found: {path}')

add_fig(doc, os.path.join(FIGURES_DIR, 'time_survival_roc_12m_ext.jpg'), 'Time-dependent AUROC for Cox, RSF, GBM survival')
add_fig(doc, os.path.join(FIGURES_DIR, 'roc_curves_binary_models_ext.jpg'), 'ROC Curves for Binary Models (Logistic Regression, Random Forest, XGBoost, MLP)')
add_fig(doc, os.path.join(FIGURES_DIR, 'calibration_1yr_models_ext.jpg'), 'Calibration Plots for Binary Models (Logistic Regression, Random Forest, XGBoost, MLP)')

doc.save(OUT_DOC)
print('External validation report saved to', OUT_DOC)
